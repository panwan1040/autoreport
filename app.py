import os
import sqlite3
from datetime import datetime, timedelta
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_file,
    flash,
    jsonify,
    abort,
)
from werkzeug.utils import secure_filename, safe_join
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image

# --- Config ---
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "uploads")
DB_PATH = os.path.join(os.path.dirname(__file__), "data.sqlite")
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "webp"}
TEMPLATE_DOCX = os.path.join(
    os.path.dirname(__file__), "4.1 รายงานประกอบการซ่อมแซมโครงสร้างอาคาร 4 ช.docx"
)

STEP_LABELS = {
    "1": "1. ก่อนการดำเนินงาน",
    "2": "2. สกัดคอนกรีตที่เสียหาย ตรวจวัดขนาดเหล็กเสริม",
    "3": "3. ขัดสนิมเหล็ก ทาน้ำยากันสนิมและน้ำยาประสานคอนกรีต",
    "4": "4. เทหุ้มด้วย NONSHRINK GROUTING MATERIAL",
    "5": "5. ทาวัสดุป้องกันผิว",
    "6": "6. ดำเนินงานแล้วเสร็จ",
}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def ensure_dirs():
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    cur = conn.cursor()
    cur.executescript(
        '''
        CREATE TABLE IF NOT EXISTS projects(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS photos(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project TEXT NOT NULL,
            point TEXT NOT NULL,
            step TEXT NOT NULL,
            filename TEXT NOT NULL,
            rel_path TEXT NOT NULL,
            uploaded_at TEXT NOT NULL
        );
        '''
    )
    conn.commit()
    conn.close()

app = Flask(__name__)
app.secret_key = os.environ.get("APP_SECRET", "devsecret")
app.config["MAX_CONTENT_LENGTH"] = 1024 * 1024 * 100  # 100MB

ensure_dirs()
init_db()

@app.template_filter("thai_step")
def thai_step(code):
    return STEP_LABELS.get(code, code)

@app.route("/", methods=["GET"])
def index():
    conn = get_db()
    projects = conn.execute("SELECT name FROM projects ORDER BY created_at DESC").fetchall()
    conn.close()
    return render_template("index.html", projects=[p["name"] for p in projects], step_labels=STEP_LABELS)

@app.route("/create_project", methods=["POST"])
def create_project():
    project = request.form.get("project_name", "").strip()
    if not project:
        flash("กรุณาระบุชื่อโปรเจ็กต์", "error")
        return redirect(url_for("index"))
    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO projects(name, created_at) VALUES (?, ?)",
            (project, datetime.now().isoformat(timespec="seconds"))
        )
        conn.commit()
        os.makedirs(os.path.join(UPLOAD_FOLDER, secure_filename(project)), exist_ok=True)
        flash(f"สร้างโปรเจ็กต์ {project} สำเร็จ", "ok")
    except sqlite3.IntegrityError:
        flash("มีชื่อโปรเจ็กต์นี้อยู่แล้ว", "warn")
    finally:
        conn.close()
    return redirect(url_for("index"))

@app.route("/upload", methods=["POST"])
def upload():
    project = request.form.get("project_select") or request.form.get("project_manual")
    if not project:
        flash("โปรดเลือกหรือระบุโปรเจ็กต์", "error")
        return redirect(url_for("index"))
    point = request.form.get("point", "").strip()
    step = request.form.get("step", "").strip()
    files = request.files.getlist("photos")
    if not point or not step:
        flash("โปรดระบุจุดและขั้นตอน", "error")
        return redirect(url_for("index"))
    if not files:
        flash("ยังไม่ได้เลือกไฟล์รูป", "error")
        return redirect(url_for("index"))
    # ensure project exists in DB
    conn = get_db()
    row = conn.execute("SELECT id FROM projects WHERE name=?", (project,)).fetchone()
    if not row:
        conn.execute("INSERT INTO projects(name, created_at) VALUES (?, ?)", (project, datetime.now().isoformat(timespec="seconds")))
        conn.commit()

    proj_folder = os.path.join(UPLOAD_FOLDER, secure_filename(project), secure_filename(point), secure_filename(step))
    os.makedirs(proj_folder, exist_ok=True)

    saved = 0
    for f in files:
        if f and allowed_file(f.filename):
            fname = secure_filename(f.filename)
            path = os.path.join(proj_folder, fname)
            f.save(path)
            rel = os.path.relpath(path, start=UPLOAD_FOLDER)
            conn.execute(
                "INSERT INTO photos(project, point, step, filename, rel_path, uploaded_at) VALUES (?, ?, ?, ?, ?, ?)",
                (project, point, step, fname, rel, datetime.now().isoformat(timespec="seconds"))
            )
            saved += 1
    conn.commit()
    conn.close()
    flash(f"อัปโหลดสำเร็จ {saved} ไฟล์ ไปยัง โปรเจ็กต์ {project} / จุด {point} / ขั้นตอน {step}", "ok")
    return redirect(url_for("index"))

@app.route("/upload_ajax", methods=["POST"])
def upload_ajax():
    """AJAX endpoint for drag and drop uploads"""
    conn = None
    try:
        project = request.form.get("project")
        point = request.form.get("point")
        step = request.form.get("step")
        files = request.files.getlist("photos")

        if not all([project, point, step]):
            return jsonify({"success": False, "message": "กรุณากรอกข้อมูลให้ครบถ้วน"})

        if not files:
            return jsonify({"success": False, "message": "ยังไม่ได้เลือกไฟล์รูป"})

        # Ensure project exists
        conn = get_db()
        row = conn.execute("SELECT id FROM projects WHERE name=?", (project,)).fetchone()
        if not row:
            conn.execute(
                "INSERT INTO projects(name, created_at) VALUES (?, ?)",
                (project, datetime.now().isoformat(timespec="seconds"))
            )
            conn.commit()

        proj_folder = os.path.join(
            UPLOAD_FOLDER,
            secure_filename(project),
            secure_filename(point),
            secure_filename(step),
        )
        os.makedirs(proj_folder, exist_ok=True)

        saved_files = []
        for f in files:
            if f and allowed_file(f.filename):
                fname = secure_filename(f.filename)
                path = os.path.join(proj_folder, fname)
                f.save(path)
                rel = os.path.relpath(path, start=UPLOAD_FOLDER)
                conn.execute(
                    "INSERT INTO photos(project, point, step, filename, rel_path, uploaded_at) VALUES (?, ?, ?, ?, ?, ?)",
                    (
                        project,
                        point,
                        step,
                        fname,
                        rel,
                        datetime.now().isoformat(timespec="seconds"),
                    ),
                )
                saved_files.append(
                    {
                        "filename": fname,
                        "size": os.path.getsize(path),
                        "url": url_for("static_file", path=rel),
                    }
                )

        conn.commit()

        return jsonify(
            {
                "success": True,
                "message": f"อัปโหลดสำเร็จ {len(saved_files)} ไฟล์",
                "files": saved_files,
            }
        )

    except Exception as e:
        return jsonify({"success": False, "message": f"เกิดข้อผิดพลาด: {str(e)}"})
    finally:
        if conn:
            conn.close()

def resize_for_docx(img_path, max_width_inches=5.5):
    # Limit width to page content area
    try:
        im = Image.open(img_path)
        # Let python-docx scale by width only; returning Inches
        return Inches(max_width_inches)
    except Exception:
        return Inches(5.5)

@app.route("/report/<project>", methods=["GET"])
def report(project):
    # Collect photos grouped by point, then step
    conn = get_db()
    rows = conn.execute(
        "SELECT id, point, step, rel_path, filename, uploaded_at FROM photos WHERE project=? ORDER BY point, step, uploaded_at",
        (project,)
    ).fetchall()
    conn.close()

    # Build nested dict
    data = {}
    for r in rows:
        point = r["point"]
        step = r["step"]
        data.setdefault(point, {}).setdefault(step, []).append(r)

    # Create document from template then clear existing content
    template = Document(TEMPLATE_DOCX)
    base_table = template.tables[0]
    doc = Document(TEMPLATE_DOCX)
    for t in doc.tables:
        t._element.getparent().remove(t._element)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Heading and report date
    doc.add_heading(
        f"รายงานประกอบการซ่อมแซมโครงสร้าง — โปรเจ็กต์ {project}", level=0
    )
    doc.add_paragraph(
        "วันที่จัดทำรายงาน: " + datetime.now().strftime("%d/%m/%Y %H:%M")
    )

    def clear_cell(cell):
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)

    cell_map = {
        "1": (0, 0),
        "2": (0, 1),
        "3": (2, 0),
        "4": (2, 1),
        "5": (4, 0),
        "6": (4, 1),
    }

    # For each point create table and insert photos
    for point in sorted(data.keys(), key=lambda x: str(x)):
        doc.add_heading(f"จุดที่ {point}", level=1)
        tbl = deepcopy(base_table._tbl)
        doc._body._element.append(tbl)
        table = doc.tables[-1]
        for step, (row, col) in cell_map.items():
            cell = table.cell(row, col)
            clear_cell(cell)
            photos = data.get(point, {}).get(step, [])
            if not photos:
                cell.add_paragraph("(ยังไม่มีรูปในขั้นตอนนี้)")
                continue
            for ph in photos:
                img_abs = os.path.join(UPLOAD_FOLDER, ph["rel_path"]).replace("\\", "/")
                try:
                    width = resize_for_docx(img_abs, 2.5)
                    par = cell.add_paragraph()
                    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = par.add_run()
                    run.add_picture(img_abs, width=width)
                    cell.add_paragraph(
                        f"ไฟล์: {ph['filename']} อัปโหลดเมื่อ {ph['uploaded_at']}"
                    ).italic = True
                except Exception as e:
                    cell.add_paragraph(
                        f"(ไม่สามารถแทรกรูป {ph['filename']}: {e})"
                    )
    # Save to temp and send
    out_dir = os.path.join(os.path.dirname(__file__), "generated")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"report_{secure_filename(project)}.docx")
    doc.save(out_path)
    return send_file(out_path, as_attachment=True, download_name=f"รายงาน_{project}.docx")

@app.route("/projects", methods=["GET"])
def list_projects():
    conn = get_db()
    rows = conn.execute("SELECT name, created_at FROM projects ORDER BY created_at DESC").fetchall()
    conn.close()
    return render_template("projects.html", projects=rows)

@app.route("/browse/<project>", methods=["GET"])  # simple gallery per project
def browse(project):
    conn = get_db()
    rows = conn.execute(
        "SELECT id, point, step, rel_path, filename, uploaded_at FROM photos WHERE project=? ORDER BY point, step, uploaded_at",
        (project,)
    ).fetchall()
    conn.close()
    items = []
    for r in rows:
        items.append({
            "id": r["id"],
            "point": r["point"],
            "step": r["step"],
            "filename": r["filename"],
            "uploaded_at": r["uploaded_at"],
            "url": url_for("static_file", path=r["rel_path"])  # we will implement a simple static serve
        })
    return render_template("browse.html", project=project, items=items, step_labels=STEP_LABELS)


@app.route("/delete_photo/<int:photo_id>", methods=["POST"])
def delete_photo(photo_id):
    conn = get_db()
    row = conn.execute("SELECT rel_path FROM photos WHERE id=?", (photo_id,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"success": False, "message": "ไม่พบรูปภาพ"})
    file_path = os.path.join(UPLOAD_FOLDER, row["rel_path"])
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
        conn.execute("DELETE FROM photos WHERE id=?", (photo_id,))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "message": str(e)})
    finally:
        conn.close()

@app.route("/uploads/<path:path>")
def static_file(path):
    """Serve user-uploaded files with basic path sanitization."""
    full_path = safe_join(UPLOAD_FOLDER, path)
    if not full_path or not os.path.isfile(full_path):
        abort(404)
    return send_file(full_path)

@app.route("/api/stats", methods=["GET"])
def get_stats():
    """Get statistics for dashboard"""
    conn = get_db()
    
    # Total projects
    total_projects = conn.execute("SELECT COUNT(*) FROM projects").fetchone()[0]
    
    # Total photos
    total_photos = conn.execute("SELECT COUNT(*) FROM photos").fetchone()[0]
    
    # Recent projects (last 7 days)
    week_ago = datetime.now() - timedelta(days=7)
    recent_projects = conn.execute(
        "SELECT COUNT(*) FROM projects WHERE created_at >= ?", 
        (week_ago.isoformat(),)
    ).fetchone()[0]
    
    conn.close()
    
    return jsonify({
        "total_projects": total_projects,
        "total_photos": total_photos,
        "recent_projects": recent_projects
    })

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
